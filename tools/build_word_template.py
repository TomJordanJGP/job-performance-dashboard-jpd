"""Build a Jobs Go Public branded Word document template.

Idempotent build script. Reads brand tokens from `theme/colors.py` (single
source of truth) and emits `templates/jgp-document-template.docx`.

Mirrors the pattern in `tools/apply_brand_to_template.py` (PowerPoint sibling):
same logo source, same DM Sans typography decisions, same hex palette.

Run: venv/bin/python tools/build_word_template.py
"""

from __future__ import annotations

import sys
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
from theme.colors import JGP_COLORS  # noqa: E402

LOGO_DIR = PROJECT_ROOT / "Brand" / "logos" / "primary"
LOGO_PNG = LOGO_DIR / "jobs-go-public-logo-full-colour.png"
LOGO_URL = "https://media.jobsgopublic.com/wp-content/uploads/2025/10/JGP-Logo_RGB.png"

OUTPUT = PROJECT_ROOT / "templates" / "jgp-document-template.docx"

DM_SANS = "DM Sans"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Brand asset acquisition
# ---------------------------------------------------------------------------

def ensure_logo() -> Path:
    """Download the full-colour PNG logo if not cached locally."""
    if LOGO_PNG.exists() and LOGO_PNG.stat().st_size > 0:
        return LOGO_PNG
    LOGO_DIR.mkdir(parents=True, exist_ok=True)
    print(f"  logo: downloading {LOGO_PNG.name}")
    with urllib.request.urlopen(LOGO_URL, timeout=30) as resp:
        LOGO_PNG.write_bytes(resp.read())
    print(f"  logo: saved {LOGO_PNG.name} ({LOGO_PNG.stat().st_size:,} B)")
    return LOGO_PNG


def logo_aspect_ratio(path: Path) -> float:
    """Width / height of the logo PNG (for accurate scaling)."""
    with Image.open(path) as img:
        return img.width / img.height


# ---------------------------------------------------------------------------
# Low-level OOXML helpers (python-docx doesn't wrap these natively)
# ---------------------------------------------------------------------------

def hex_to_rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value.lstrip("#").upper())


def set_run_font(run, *, name: str = DM_SANS, size_pt: int | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 colour_hex: str | None = None) -> None:
    """Configure a run with DM Sans by default, applying east-asian + complex-
    script fallbacks the same way the PPTX builder writes <a:latin>/<a:ea>/<a:cs>."""
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if colour_hex is not None:
        run.font.color.rgb = hex_to_rgb(colour_hex)


def set_paragraph_shading(paragraph, hex_value: str) -> None:
    """Apply a solid background fill to a paragraph (<w:shd>)."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_value.lstrip("#").upper())
    ppr.append(shd)


def set_paragraph_border(paragraph, *, top: bool = False, sides: bool = False,
                          colour_hex: str = "cbb9df", size_eighths: int = 8) -> None:
    """Add borders to a paragraph. size_eighths is in 1/8-of-a-point units."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    sides_to_set = []
    if top:
        sides_to_set.append("top")
    if sides:
        sides_to_set.extend(["left", "right", "bottom"])
    for side in sides_to_set:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size_eighths))
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), colour_hex.lstrip("#").upper())
        pbdr.append(border)


def set_cell_shading(cell, hex_value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_value.lstrip("#").upper())
    tc_pr.append(shd)


def add_field(paragraph, instr_text: str, *, run_kwargs: dict | None = None):
    """Insert a Word field (e.g. DATE, PAGE) into a paragraph."""
    run_kwargs = run_kwargs or {}
    run = paragraph.add_run()
    set_run_font(run, **run_kwargs)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr_text)
    # Provide a fallback display value so the field shows something before
    # Word recomputes it. Word refreshes on open.
    placeholder = OxmlElement("w:r")
    placeholder_rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), DM_SANS)
    placeholder_rpr.append(rfonts)
    placeholder.append(placeholder_rpr)
    txt = OxmlElement("w:t")
    txt.text = "—"
    placeholder.append(txt)
    fld.append(placeholder)
    run._r.addnext(fld)
    return run


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def configure_styles(document: Document) -> None:
    """Define every paragraph style used throughout the template.

    Existing built-in styles (Title, Heading 1/2/3, Normal, Caption) are
    re-configured rather than re-created — Word treats those as well-known.
    Custom styles ("Body", "JGP Callout", "JGP Eyebrow") are added fresh.
    """
    styles = document.styles

    # --- Normal (default body) ---
    normal = styles["Normal"]
    normal.font.name = DM_SANS
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_to_rgb(JGP_COLORS["text_primary"])
    # rFonts fallbacks so non-Latin chars also pick up DM Sans
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), DM_SANS)
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_after = Pt(8)

    # --- Title (cover) ---
    title = styles["Title"]
    title.font.name = DM_SANS
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = hex_to_rgb(JGP_COLORS["deep_blue"])
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    # --- Subtitle (cover) ---
    subtitle = styles["Subtitle"]
    subtitle.font.name = DM_SANS
    subtitle.font.size = Pt(16)
    subtitle.font.bold = False
    subtitle.font.color.rgb = hex_to_rgb(JGP_COLORS["supporting"])
    subtitle.font.italic = False
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(36)

    # --- Heading 1 ---
    h1 = styles["Heading 1"]
    h1.font.name = DM_SANS
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = hex_to_rgb(JGP_COLORS["deep_blue"])
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # --- Heading 2 ---
    h2 = styles["Heading 2"]
    h2.font.name = DM_SANS
    h2.font.size = Pt(16)
    h2.font.bold = False
    h2.font.color.rgb = hex_to_rgb(JGP_COLORS["primary"])
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # --- Heading 3 ---
    h3 = styles["Heading 3"]
    h3.font.name = DM_SANS
    h3.font.size = Pt(13)
    h3.font.bold = False
    h3.font.color.rgb = hex_to_rgb(JGP_COLORS["deep_blue"])
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # --- Caption ---
    caption = styles["Caption"]
    caption.font.name = DM_SANS
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = hex_to_rgb(JGP_COLORS["text_secondary"])
    caption.paragraph_format.space_after = Pt(12)

    # --- JGP Eyebrow (custom — small uppercase tag above section headings) ---
    from docx.enum.style import WD_STYLE_TYPE  # local to keep top imports clean

    if "JGP Eyebrow" not in [s.name for s in styles]:
        eyebrow = styles.add_style("JGP Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
        eyebrow.base_style = styles["Normal"]
        eyebrow.font.name = DM_SANS
        eyebrow.font.size = Pt(9)
        eyebrow.font.bold = True
        eyebrow.font.color.rgb = hex_to_rgb(JGP_COLORS["primary"])
        eyebrow.font.all_caps = True
        eyebrow.paragraph_format.space_before = Pt(18)
        eyebrow.paragraph_format.space_after = Pt(0)
        eyebrow.paragraph_format.keep_with_next = True

    # --- JGP Callout (custom — beige panel for highlighted notes) ---
    if "JGP Callout" not in [s.name for s in styles]:
        callout = styles.add_style("JGP Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.base_style = styles["Normal"]
        callout.font.name = DM_SANS
        callout.font.size = Pt(11)
        callout.font.color.rgb = hex_to_rgb(JGP_COLORS["deep_blue"])
        callout.paragraph_format.space_before = Pt(12)
        callout.paragraph_format.space_after = Pt(12)
        callout.paragraph_format.left_indent = Cm(0.3)
        callout.paragraph_format.right_indent = Cm(0.3)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def configure_page(document: Document) -> None:
    """A4 page size, 2.5 cm margins all round."""
    for section in document.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)


# ---------------------------------------------------------------------------
# Header & footer
# ---------------------------------------------------------------------------

def configure_header_footer(document: Document) -> None:
    section = document.sections[0]
    # Different first page so the cover stays clean.
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)

    # --- Running header (page 2+) ---
    header = section.header
    header_para = header.paragraphs[0]
    header_para.style = document.styles["Caption"]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("Jobs Go Public  ·  Document title")
    set_run_font(run, size_pt=9, italic=True,
                 colour_hex=JGP_COLORS["text_secondary"])

    # --- First-page header (cover) — leave empty for a clean cover ---
    fp_header = section.first_page_header
    fp_header.paragraphs[0].text = ""

    # --- Running footer (page 2+) ---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.style = document.styles["Caption"]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_border(
        footer_para, top=True,
        colour_hex=JGP_COLORS["border"], size_eighths=8,
    )
    label_run = footer_para.add_run("Page ")
    set_run_font(label_run, size_pt=9, italic=False,
                 colour_hex=JGP_COLORS["text_secondary"])
    add_field(footer_para, "PAGE \\* MERGEFORMAT",
              run_kwargs={"size_pt": 9, "colour_hex": JGP_COLORS["deep_blue"],
                          "bold": True})
    of_run = footer_para.add_run(" of ")
    set_run_font(of_run, size_pt=9, italic=False,
                 colour_hex=JGP_COLORS["text_secondary"])
    add_field(footer_para, "NUMPAGES \\* MERGEFORMAT",
              run_kwargs={"size_pt": 9, "colour_hex": JGP_COLORS["deep_blue"],
                          "bold": True})

    # --- First-page footer (cover) — empty ---
    fp_footer = section.first_page_footer
    fp_footer.paragraphs[0].text = ""


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def build_cover(document: Document, logo_path: Path) -> None:
    """Logo top, big title, subtitle, date, then page break."""
    # Logo
    logo_para = document.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_before = Pt(36)
    logo_para.paragraph_format.space_after = Pt(48)
    logo_run = logo_para.add_run()
    aspect = logo_aspect_ratio(logo_path)
    logo_width_cm = 4.5
    logo_run.add_picture(str(logo_path),
                         width=Cm(logo_width_cm),
                         height=Cm(logo_width_cm / aspect))

    # Title
    title_para = document.add_paragraph(style="Title")
    title_run = title_para.add_run("Document title")
    set_run_font(title_run, size_pt=32, bold=True,
                 colour_hex=JGP_COLORS["deep_blue"])

    # Subtitle
    subtitle_para = document.add_paragraph(style="Subtitle")
    subtitle_run = subtitle_para.add_run("Subtitle or section")
    set_run_font(subtitle_run, size_pt=16,
                 colour_hex=JGP_COLORS["supporting"])

    # Metadata block: prepared / date — uses a 2-col table so labels and
    # values align cleanly without tab-stops.
    meta_table = document.add_table(rows=2, cols=2)
    meta_table.autofit = False
    meta_table.columns[0].width = Cm(4.0)
    meta_table.columns[1].width = Cm(11.5)
    for row in meta_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(4)

    label_run = meta_table.cell(0, 0).paragraphs[0].add_run("Prepared")
    set_run_font(label_run, size_pt=9, bold=True,
                 colour_hex=JGP_COLORS["primary"])
    # All-caps via run-level OOXML (caption built-in already has all_caps;
    # easier to just write the letters in caps and apply small weight).
    val_run = meta_table.cell(0, 1).paragraphs[0].add_run(
        "Author name  ·  Team")
    set_run_font(val_run, size_pt=11,
                 colour_hex=JGP_COLORS["text_primary"])

    date_label_run = meta_table.cell(1, 0).paragraphs[0].add_run("Date")
    set_run_font(date_label_run, size_pt=9, bold=True,
                 colour_hex=JGP_COLORS["primary"])
    date_value_para = meta_table.cell(1, 1).paragraphs[0]
    add_field(date_value_para, 'DATE \\@ "d MMMM yyyy" \\* MERGEFORMAT',
              run_kwargs={"size_pt": 11,
                          "colour_hex": JGP_COLORS["text_primary"]})

    # Page break to start body content on page 2
    document.add_page_break()


# ---------------------------------------------------------------------------
# Sample body content (style showcase)
# ---------------------------------------------------------------------------

def build_body(document: Document) -> None:
    """Add one of every style so the user can see/copy each in Word."""

    # --- Eyebrow + Heading 1 ---
    eyebrow_para = document.add_paragraph("Section 01", style="JGP Eyebrow")
    document.add_paragraph(
        "Heading 1 — the page-level title", style="Heading 1"
    )
    document.add_paragraph(
        "Body copy uses DM Sans Regular at 11 pt with 1.4 line spacing. "
        "Text colour is deep-blue (#240f45), the primary text token in the "
        "brand kit. Keep paragraphs short — three to four sentences typically "
        "reads best at this size. Use this style for everything that isn't "
        "explicitly a heading, caption, or callout."
    )

    document.add_paragraph(
        "Heading 2 — section title", style="Heading 2"
    )
    document.add_paragraph(
        "Another body paragraph follows so you can see how the spacing "
        "between Heading 2 and body copy lands on the page. Tighter than "
        "Heading 1 above, but still gives the eye somewhere to rest."
    )

    document.add_paragraph(
        "Heading 3 — sub-section", style="Heading 3"
    )
    document.add_paragraph(
        "Use Heading 3 for nested structure inside a long section. It picks "
        "up the deep-blue brand colour at 13 pt — visually quieter than the "
        "purple Heading 2 but still clearly a heading."
    )

    # --- Bulleted list ---
    document.add_paragraph(
        "A bulleted list", style="Heading 3"
    )
    for item in (
        "Bullets use the standard List Bullet style.",
        "DM Sans body settings carry through.",
        "Keep items parallel in structure — verb-first reads cleanest.",
    ):
        bullet = document.add_paragraph(item, style="List Bullet")
        for run in bullet.runs:
            set_run_font(run, size_pt=11,
                         colour_hex=JGP_COLORS["text_primary"])

    # --- Numbered list ---
    document.add_paragraph(
        "A numbered list", style="Heading 3"
    )
    for item in (
        "First step in a sequence.",
        "Second step — Word handles the numbering.",
        "Third step closes it out.",
    ):
        numbered = document.add_paragraph(item, style="List Number")
        for run in numbered.runs:
            set_run_font(run, size_pt=11,
                         colour_hex=JGP_COLORS["text_primary"])

    # --- Beige callout ---
    document.add_paragraph(
        "A callout panel", style="Heading 3"
    )
    callout = document.add_paragraph(style="JGP Callout")
    callout_run = callout.add_run(
        "Use the JGP Callout style for a highlighted note, takeaway, or "
        "summary. The beige panel (#f0f3e1) ties back to the brand kit's "
        "chart-supporting background."
    )
    set_run_font(callout_run, size_pt=11,
                 colour_hex=JGP_COLORS["deep_blue"])
    set_paragraph_shading(callout, JGP_COLORS["beige"])
    set_paragraph_border(callout, sides=True,
                          colour_hex=JGP_COLORS["light_purple"],
                          size_eighths=4)

    # --- Branded table ---
    document.add_paragraph(
        "A branded table", style="Heading 3"
    )
    table = document.add_table(rows=4, cols=3)
    table.autofit = False
    col_widths = (Cm(5.5), Cm(5.0), Cm(5.0))
    for idx, width in enumerate(col_widths):
        table.columns[idx].width = width
        for cell in table.columns[idx].cells:
            cell.width = width

    headers = ("Metric", "Value", "Change")
    rows = (
        ("Total applies", "12,480", "+18%"),
        ("Cost per apply", "£2.41", "−9%"),
        ("Conversion rate", "4.3%", "+0.6 pp"),
    )

    # Header row
    for col_idx, label in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_shading(cell, JGP_COLORS["primary"])
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        run = cell.paragraphs[0].add_run(label)
        set_run_font(run, size_pt=10, bold=True,
                     colour_hex=JGP_COLORS["white"])

    # Body rows with alternating beige stripes
    for row_idx, row_values in enumerate(rows, start=1):
        is_alt = (row_idx % 2 == 1)
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            if is_alt:
                set_cell_shading(cell, JGP_COLORS["beige"])
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size_pt=10,
                         colour_hex=JGP_COLORS["text_primary"])

    caption_para = document.add_paragraph(
        "Caption — small italic descriptor below figures, tables or "
        "screenshots. Uses secondary-text colour (#4f4360).",
        style="Caption",
    )
    # Ensure caption font runs are DM Sans even though the style says so
    # (Word can otherwise inherit Cambria from the underlying style).
    for run in caption_para.runs:
        set_run_font(run, size_pt=9, italic=True,
                     colour_hex=JGP_COLORS["text_secondary"])


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    print("Building JGP Word template")
    print(f"  output: {OUTPUT.relative_to(PROJECT_ROOT)}")

    logo_path = ensure_logo()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_page(document)
    configure_styles(document)
    configure_header_footer(document)
    build_cover(document, logo_path)
    build_body(document)

    document.save(OUTPUT)
    print(f"  saved: {OUTPUT.name} ({OUTPUT.stat().st_size:,} B)")
    print("Done.")


if __name__ == "__main__":
    main()
